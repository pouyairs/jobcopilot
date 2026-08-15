from io import BytesIO

from pypdf import PdfReader
from docx import Document


def extract_pdf_text(uploaded_file):

    uploaded_file.seek(0)

    reader = PdfReader(uploaded_file)

    parts = []

    for page in reader.pages:

        text = page.extract_text()

        if text:
            parts.append(text)

    return "\n".join(parts).strip()


def extract_docx_text(uploaded_file):

    uploaded_file.seek(0)

    file_bytes = BytesIO(
        uploaded_file.read()
    )

    document = Document(file_bytes)

    parts = []

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:
            parts.append(text)

    # Also read tables
    for table in document.tables:

        for row in table.rows:

            cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if cells:
                parts.append(
                    " | ".join(cells)
                )

    return "\n".join(parts).strip()


def extract_cv_text(uploaded_file):

    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):

        text = extract_pdf_text(
            uploaded_file
        )

    elif filename.endswith(".docx"):

        text = extract_docx_text(
            uploaded_file
        )

    else:

        raise ValueError(
            "Unsupported file format."
        )

    if len(text.strip()) < 100:

        raise ValueError(
            "Not enough readable text was found. "
            "Please upload a text-based PDF or DOCX file."
        )

    return text